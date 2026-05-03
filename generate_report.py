from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── styles ──

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_heading(paragraph, level):
    sizes = {1: 20, 2: 15, 3: 13}
    paragraph.style = doc.styles[f"Heading {level}"]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(sizes.get(level, 12))

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(text="", bold=False, italic=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def add_code_block(text):
    """Add a monospaced code-style paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
    # data rows
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    doc.add_paragraph()  # spacing after table

# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Crypto4")
run.bold = True
run.font.size = Pt(28)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Project Report")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_p.add_run("Security of Computers — Password Manager Project").italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════

add_heading("1. Project Overview")
add_paragraph(
    "Crypto4 is a password manager built as a security-of-computers coursework project. "
    "It implements four cryptographic modules from scratch (hashing, symmetric encryption, "
    "digital signatures, and key exchange), integrates them into a functioning vault application, "
    "and applies them to six CTF challenges. The application ships with both a CLI prototype "
    "(main.py) and a Streamlit web interface (app.py)."
)

# ══════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════

add_heading("2. Architecture")

# 2.1
add_heading("2.1 Layer Separation", level=2)
add_paragraph(
    "The codebase is split into three clear tiers:"
)
add_table(
    ["Layer", "Files", "Responsibility"],
    [
        ["Crypto primitives", "modules/", "Pure crypto, no UI dependencies"],
        ["Business logic", "vault_service.py, diffie_hellman_export.py", "Vault operations, DH protocol"],
        ["Presentation", "UI/", "Streamlit pages and components"],
    ]
)
add_paragraph(
    "vault_service.py carries a module-level docstring explicitly stating "
    '"No Streamlit imports — purely data operations." '
    "This boundary was a deliberate design decision that kept the crypto logic testable "
    "and reusable independently of the UI framework."
)

# 2.2
add_heading("2.2 Per-User Vault Layout", level=2)
add_paragraph("Each registered user gets an isolated directory:")
add_code_block(
    "vaults/\n"
    "  {username}/\n"
    "    vault.json          ← encrypted vault + ElGamal signature\n"
    "    keys/\n"
    "      {username}_private.json\n"
    "      {username}_public.json"
)
add_paragraph("The vault file format is a two-field JSON envelope:")
add_code_block(
    '{\n'
    '  "encrypted_vault": "<base64 AES-GCM ciphertext>",\n'
    '  "signature": { "r": "<hex>", "s": "<hex>" }\n'
    '}'
)
add_paragraph(
    "The signature covers the encrypted ciphertext (encrypt-then-sign), so tampering with "
    "the ciphertext is detectable before any decryption attempt."
)

# 2.3
add_heading("2.3 Two-Interface Design", level=2)
add_paragraph(
    "main.py was the initial CLI prototype for testing each module in isolation "
    "(hash a file, encrypt/decrypt with AES or RSA, check password strength, manage ElGamal keys). "
    "The Streamlit app (app.py) is the final product. The CLI was kept and is still functional, "
    "making it useful for demonstration and debugging."
)

# ══════════════════════════════════════════════════════════════════
# 3. ALGORITHM CHOICES
# ══════════════════════════════════════════════════════════════════

add_heading("3. Algorithm Choices")

# 3.1
add_heading("3.1 AES-256-GCM — Vault Encryption", level=2)
add_paragraph(
    "AES-GCM was chosen for the vault because it provides authenticated encryption: "
    "confidentiality and integrity in a single pass. Any modification to the ciphertext makes "
    "decryption fail, providing a second integrity layer on top of the ElGamal signature."
)
add_paragraph("Implementation details (modules/encryption.py):")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Library: ").bold = True
p.add_run("PyCryptodomex (Cryptodome.Cipher.AES)")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Mode: ").bold = True
p.add_run("GCM with a 12-byte random nonce per encryption")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Key derivation: ").bold = True
p.add_run("SHA-256(master_password) → 32-byte AES key")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Ciphertext format: ").bold = True
p.add_run("nonce ‖ ciphertext ‖ tag, base64-encoded")
doc.add_paragraph()
add_paragraph(
    "The same AES_Encryption class is reused for the DH vault transfer — the session key "
    "derived from the DH shared secret is fed into it, eliminating duplicate encryption code. "
    "A get_secret_key method using PBKDF2-HMAC exists in the class but is not used in "
    "production. Only the SHA-256 derivation path is active (noted as 'per spec' in the code)."
)

# 3.2
add_heading("3.2 ElGamal Digital Signatures — Vault Integrity", level=2)
add_paragraph(
    "ElGamal signatures were implemented from scratch rather than calling a library "
    "(modules/sign.py, modules/verify.py)."
)
add_paragraph("The signing equation:")
add_code_block(
    "H  = SHA-256(message) mod (p-1)\n"
    "r  = alpha^k mod p\n"
    "s  = k_inv * (H - x*r) mod (p-1)"
)
add_paragraph("Verification checks:")
add_code_block("alpha^H ≡ y^r * r^s  (mod p)")
add_paragraph("Key uses of ElGamal signatures in the system:")
for item in [
    "Signing the vault file after every write operation",
    "Signing ephemeral DH public keys during vault export to prevent MITM substitution",
    "Signing the re-encrypted vault payload before transmission",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph()

# 3.3
add_heading("3.3 Diffie-Hellman — Secure Vault Transfer", level=2)
add_paragraph(
    "The vault export/import flow (diffie_hellman_export.py) implements a full authenticated "
    "DH exchange:"
)
steps = [
    "Device 1 generates an ephemeral DH keypair, signs the public value with its ElGamal private key, sends the signed package to Device 2.",
    "Device 2 verifies D1's signature (aborts on failure), generates its own ephemeral DH keypair, signs it, sends back.",
    "Device 1 verifies D2's signature, computes shared secret → derives a 256-bit session key via SHA-256(shared_secret).",
    "Device 1 decrypts the vault with its master password, re-encrypts with the session key, signs the result, serialises to a JSON bundle.",
    "Device 2 re-derives the same shared secret (using D1's DH public key with its own DH private key), verifies D1's signature on the bundle, decrypts, re-encrypts under its own master password, saves and re-signs with its ElGamal key.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)
doc.add_paragraph()
add_paragraph(
    "Signing the DH public keys before exchange was a deliberate choice to prevent a "
    "man-in-the-middle from swapping the ephemeral keys."
)

# 3.4
add_heading("3.4 Shared Cryptographic Parameters", level=2)
add_paragraph(
    "Parameters for both ElGamal and DH are generated dynamically at startup (modules/config.py):"
)
for item in [
    "A 64-bit safe prime p = 2q + 1 is generated using sympy.isprime",
    "A generator alpha of the order-q subgroup is found by random sampling (avoiding the slow sequential search)",
    "Both ELGAMAL_PARAMS and DH_PARAMS share the same p and alpha",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph()
add_paragraph(
    "The commented-out small-prime block (p=23, alpha=5) in config.py shows the development "
    "path: hardcoded toy parameters first, then replaced with generated safe primes."
)

# 3.5
add_heading("3.5 Password Hashing — bcrypt + zxcvbn", level=2)
add_paragraph(
    "bcrypt is used for password hashing in the password.py module (demonstrated via CLI). "
    "The salt is generated per-call by bcrypt.gensalt(). "
    "zxcvbn provides entropy-based strength estimation. A minimum score of 3 is enforced at "
    "account registration and vault import (when setting the new master password). "
    "This prevents users from protecting their vault with easily guessable passwords."
)

# 3.6
add_heading("3.6 SHA-256 — Pervasive Hash Function", level=2)
add_paragraph("SHA-256 appears at every level of the stack:")
add_table(
    ["Use", "Location"],
    [
        ["File integrity checking", "modules/hash.py"],
        ["Message hash for ElGamal signature", "modules/sign.py, modules/verify.py"],
        ["AES key derivation from master password", "modules/encryption.py"],
        ["DH session key derivation", "diffie_hellman_export.py"],
    ]
)

# 3.7
add_heading("3.7 RSA-2048 with OAEP (Demo Only)", level=2)
add_paragraph(
    "RSA encryption/decryption is implemented in modules/encryption.py using "
    "cryptography.hazmat with OAEP padding and SHA-256. It is accessible through the CLI "
    "for demonstration but is not used in the vault application, which relies on AES for "
    "encryption and ElGamal for signatures."
)

# ══════════════════════════════════════════════════════════════════
# 4. WORK INTEGRATION CHALLENGES
# ══════════════════════════════════════════════════════════════════

add_heading("4. Work Integration Challenges")

# 4.1
add_heading("4.1 Module 3 (Signing) Was a Blocker for Module 2 (Vault Encryption)", level=2)
add_paragraph(
    "The clearest evidence of integration friction appears as comments in "
    "modules/vault_encryption.py:"
)
add_code_block(
    "def add(self, ...):\n"
    "    ...\n"
    "    # 6) Resign vault file\n"
    "    # Module 3 function -> Waiting\n"
    "    return self._save_vault(entries)"
)
add_paragraph(
    "The same comment appears in update() and delete(). The vault CRUD operations were built "
    "while the ElGamal signing module was still in progress, so placeholder comments were "
    "left marking where the signing calls needed to be wired in once Module 3 was delivered. "
    "This is a textbook dependency ordering problem in team development — the encryption team "
    "could not fully close out their module until the signature team shipped theirs."
)

# 4.2
add_heading("4.2 Two Parallel Vault Abstractions", level=2)
add_paragraph("There are two separate vault management implementations that coexist:")
for item in [
    "modules/vault_encryption.py — a VaultEncryption class with add(), retrieve(), update(), delete() methods",
    "vault_service.py — standalone functions (load_entries, save_entries, etc.)",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph()
add_paragraph(
    "The running application uses vault_service.py. The VaultEncryption class appears to be "
    "an earlier iteration built when the sign/verify integration was still pending. When the "
    "team consolidated, a new service layer was written rather than retrofitting the old class. "
    "Both are present in the final codebase."
)

# 4.3
add_heading("4.3 Dual Encryption Libraries", level=2)
add_paragraph(
    "Two different encryption libraries are imported in modules/encryption.py:"
)
for item in [
    "cryptography.hazmat (from the cryptography package) — used by the aes_ed and rsa_ed standalone functions",
    "Cryptodome.Cipher.AES (from pycryptodomex) — used by the AES_Encryption class",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph()
add_paragraph(
    "This reflects different team members working from different references before converging. "
    "The AES_Encryption class became the canonical production path; the aes_ed function "
    "remains as a CLI demo."
)

# 4.4
add_heading("4.4 Lazy Imports to Avoid Circular Dependencies", level=2)
add_paragraph(
    "In vault_service.py, the diffie_hellman_export module is imported inside the function "
    "body rather than at the top of the file:"
)
add_code_block(
    "def build_export_bundle(...):\n"
    "    from diffie_hellman_export import (\n"
    "        device1_start_exchange,\n"
    "        ...\n"
    "    )"
)
add_paragraph(
    "This is a sign that when the DH module was integrated into the service layer, a circular "
    "import was encountered and resolved by deferring the import. It works correctly but is a "
    "structural seam left by the integration."
)

# 4.5
add_heading("4.5 Unused PBKDF2 Method", level=2)
add_paragraph(
    "The AES_Encryption class contains a get_secret_key method that implements PBKDF2-HMAC "
    "with 65,535 iterations, which is the stronger approach for key derivation. The constructor, "
    "however, uses plain SHA-256. This suggests PBKDF2 was prototyped (or came from a different "
    "team member's approach) but was not adopted when the spec called for SHA-256 as the key "
    "derivation path. The method sits unused in production."
)

# ══════════════════════════════════════════════════════════════════
# 5. CTF CHALLENGES
# ══════════════════════════════════════════════════════════════════

add_heading("5. CTF Challenges")
add_paragraph(
    "Six CTF challenges were solved as part of the project:"
)
add_table(
    ["#", "Title", "Technique"],
    [
        ["CTF 1", "Packet Analysis",
         "Manual pcap parsing (no libraries), TCP stream reassembly, base64 decode from port 4444"],
        ["CTF 2", "Visual Cryptography",
         "XOR two PNG layers pixel-by-pixel using numpy to cancel one-time-pad noise and reveal hidden image"],
        ["CTF 3", "Bit Shift Cipher",
         "Right-shift each number in a file by 1 bit to recover ASCII character values"],
        ["CTF 4", "LSB Steganography",
         "Extract least-significant bits from grayscale PNG pixels, reconstruct 8-bit bytes to read hidden message"],
        ["CTF 5", "CBC Padding Oracle",
         "Byte-by-byte decryption of AES-CBC ciphertext via a remote padding oracle, with false-positive check at last byte"],
        ["CTF 6", "Weak RSA Factorisation",
         "Factor small RSA modulus with sympy.factorint, recover d = e^-1 mod φ(n), decrypt ciphertext"],
    ]
)
add_paragraph(
    "CTF 5 is the most technically demanding — it implements the full padding oracle attack, "
    "including the edge-case disambiguation step for the last byte (where a padding value of "
    "\\x01 can be a false positive when the preceding byte accidentally produces valid padding)."
)

# ══════════════════════════════════════════════════════════════════
# 6. DEPENDENCIES
# ══════════════════════════════════════════════════════════════════

add_heading("6. Dependencies")
add_table(
    ["Package", "Role"],
    [
        ["pycryptodomex", "AES-GCM encryption (vault)"],
        ["cryptography",  "AES-GCM (CLI demo), RSA-OAEP"],
        ["sympy",         "Safe prime generation, RSA factorisation (CTF 6)"],
        ["bcrypt",        "Password hashing"],
        ["zxcvbn",        "Password strength estimation"],
        ["streamlit",     "Web UI"],
        ["pillow",        "Image manipulation (CTF 2, CTF 4)"],
        ["numpy",         "Pixel-level XOR operations (CTF 2)"],
        ["requests",      "CBC padding oracle HTTP calls (CTF 5)"],
    ]
)

# ── save ──
out = "Crypto4_Report.docx"
doc.save(out)
print(f"Saved: {out}")
